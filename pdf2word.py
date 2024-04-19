import PyPDF2
import tabula
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

def pdf_to_word(pdf_path, word_path=None):
    # 1. 使用PyPDF2读取PDF文件
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)

        # 2. 创建一个Word文档对象
        if not word_path:
            word_path = pdf_path.replace('.pdf', '.docx')
        doc = Document()

        # 3. 遍历PDF页面并将其内容添加到Word文档中
        for page_num in range(len(reader.pages)):
            page_obj = reader.pages[page_num]
            text = page_obj.extract_text()
            doc.add_paragraph(text)

        # 4. 保存Word文档
        doc.save(word_path)

        print(f"PDF file '{pdf_path}' successfully converted to Word document '{word_path}'.")

def pdf_to_word_tabula(pdf_path, word_path=None):
    # 从PDF中读取表格
    tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)

    # 创建一个新的Word文档
    doc = Document()

    # 遍历从PDF中提取的所有表格
    for i, table in enumerate(tables):
        # 添加一个新的表格到Word文档中
        word_table = doc.add_table(rows=len(table) + 1, cols=len(table.columns))

        # 设置表格的标题（如果有的话）
        if i == 0:
            word_table.style = 'Table Grid'  # 可以根据需要更改表格样式
            hdr_cells = word_table.rows[0].cells
            for j, col_name in enumerate(table.columns):
                hdr_cells[j].text = col_name

        # 填充表格的其余部分
        for row_idx, row_data in table.iterrows():
            row_cells = word_table.rows[row_idx + 1].cells
            for col_idx, col_data in enumerate(row_data):
                row_cells[col_idx].text = str(col_data)

        # 添加一个空行，以便下一个表格有一个新的开始
        doc.add_paragraph()

    # 保存Word文档
    doc.save(word_path)

    print(f"PDF file '{pdf_path}' successfully converted to Word document '{word_path}'.")

# 使用示例
pdf_to_word_tabula('《奇妙的想象》教学设计——吕玲.pdf', '《奇妙的想象》教学设计.docx')


# 使用示例
#pdf_to_word('《奇妙的想象》教学设计——吕玲.pdf')
